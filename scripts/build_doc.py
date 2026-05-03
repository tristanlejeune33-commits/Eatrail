#!/usr/bin/env python3
"""Build Savora / Mama Map startup concept document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
PRIMARY = RGBColor(0xC8, 0x5A, 0x3A)   # terracotta
DARK    = RGBColor(0x1C, 0x1B, 0x1A)
MUTED   = RGBColor(0x6B, 0x65, 0x60)
ACCENT  = RGBColor(0x3E, 0x5C, 0x3E)
GOLD    = RGBColor(0xD9, 0xA4, 0x41)
CREAM   = "F5EEE0"
POP_RED = RGBColor(0xFF, 0x4E, 0x3A)

doc = Document()

# ---- Global style defaults ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

# ---- Page setup: US Letter with 1" margins ----
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ---- Heading styles ----
def style_heading(level, size, color=PRIMARY, space_before=18, space_after=10):
    s = doc.styles[f'Heading {level}']
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    s.font.name = 'Calibri'
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)

style_heading(1, 22, PRIMARY, 22, 10)
style_heading(2, 16, DARK, 14, 6)
style_heading(3, 13, DARK, 10, 4)

# ---- Helpers ----
def add_header_footer():
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        run = hp.add_run("Savora / Mama Map")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        # tab + right-aligned context
        tab = hp.add_run("\tConcept produit · NYC · 2026")
        tab.font.size = Pt(9)
        tab.font.color.rgb = MUTED
        # tab stop right
        tab_stops = hp.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), alignment=2)  # RIGHT=2
        # Footer with page number
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        r1 = fp.add_run("Cook the world. Shop next door.")
        r1.italic = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = MUTED
        r2 = fp.add_run("\tPage ")
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED
        # Insert PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r3 = fp.add_run()
        r3.font.size = Pt(9)
        r3.font.color.rgb = MUTED
        r3._r.append(fldChar1)
        r3._r.append(instrText)
        r3._r.append(fldChar2)
        ts = fp.paragraph_format.tab_stops
        ts.add_tab_stop(Inches(6.5), alignment=2)

def set_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def h1(text, color=PRIMARY):
    p = doc.add_heading('', level=1)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = True
    return p

def h2(text):
    p = doc.add_heading('', level=2)
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.bold = True
    return p

def h3(text):
    p = doc.add_heading('', level=3)
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.bold = True
    return p

def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def parse_markdown_bold(paragraph, text):
    """Split text by **bold** and add runs accordingly."""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def para(text, italic=False, color=None, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    parse_markdown_bold(p, text)
    for run in p.runs:
        if italic:
            run.italic = True
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def lead(text):
    return para(text, italic=True, color=MUTED, size=12, space_after=10)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    parse_markdown_bold(p, text)
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p

def num_item(text):
    p = doc.add_paragraph(style='List Number')
    parse_markdown_bold(p, text)
    p.paragraph_format.space_after = Pt(3)
    return p

def quote(text, color=PRIMARY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\u201C{text}\u201D')
    run.italic = True
    run.font.color.rgb = color
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def divider():
    p = doc.add_paragraph()
    # bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C85A3A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    # remove any existing tcBorders to avoid duplicates
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def simple_table(headers, rows, col_widths_in, header_color="C85A3A"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    # set column widths
    for i, w in enumerate(col_widths_in):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    # header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_border(hdr[i])
        set_shading(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.text = ""
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    # rows
    for ridx, row in enumerate(rows):
        cells = t.rows[ridx + 1].cells
        fill = "FFFFFF" if ridx % 2 == 0 else "FAF7F2"
        for i, val in enumerate(row):
            set_cell_border(cells[i])
            set_shading(cells[i], fill)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[i].paragraphs[0]
            p.text = ""
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
    # spacing after
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)
    return t

# =====================================================
# COVER
# =====================================================
# Spacer
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SAVORA / MAMA MAP")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cook the world. Shop next door.")
r.italic = True
r.font.size = Pt(18)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(40)

for text in [
    "Concept produit · Branding · UX · MVP · Go-to-market",
    "Deux versions — Startup sérieuse + Agressive croissance / viralité",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NYC — 2026")
r.bold = True
r.font.color.rgb = GOLD
r.font.size = Pt(11)
p.paragraph_format.space_before = Pt(18)

page_break()

# =====================================================
# INTRO
# =====================================================
h1("Intro — le produit en une phrase")
lead("Tu veux cuisiner un plat du monde → l'app te donne les ingrédients → et te dit où les acheter près de chez toi, au meilleur prix et le plus authentique possible.")
para("Trois douleurs réelles à New York : (1) trouver la bonne recette authentique (pas une version fade d'un blog food US), (2) savoir où acheter les ingrédients rares (piment gochugaru, sauce nuoc mam premium, maseca, tamarin frais), (3) ne pas faire 4 quartiers pour 3 courses. Savora / Mama Map résout les trois d'un coup.")
para("Marché : NYC, 8,3M d'habitants dont ~37% nés hors des US, ~800 langues parlées, ~24k commerces alimentaires spécialisés. Une capitale mondiale de la food où personne n'a encore connecté recette + carte + authenticité dans une seule expérience mobile. Scalable ensuite sur LA, SF, Miami, Toronto, Londres, Paris, Berlin, Singapour.")
divider()
para("**Ce document présente DEUX versions stratégiques du même produit :**")
bullet("**Version 1 — Startup sérieuse** : positionnement premium accessible, investor-ready, branding sobre, modèle de revenus diversifié.")
bullet("**Version 2 — Agressive croissance / viralité** : positionnement pop & social-first, mécaniques virales, growth loops, content-as-product.")
para("Les deux partent du même cœur produit mais optimisent des KPIs différents : rétention + LTV vs. vélocité d'acquisition + K-factor.")
page_break()

# =====================================================
# VERSION 1 — STARTUP SÉRIEUSE
# =====================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VERSION 1")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MUTED
p.paragraph_format.space_before = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Startup sérieuse")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Premium accessible · culture · food · tech")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(18)

divider()

# 1. BRANDING
h1("1. Branding")
h2("Nom de marque — 3 propositions")
para("**SAVORA** (recommandé)")
bullet("« Savor » + terminaison douce en « a » — racines romanes universelles.")
bullet("3 syllabes, 6 lettres, phonétique stable en EN / FR / ES / IT.")
bullet("Disponible en .app / .co, handles @savora libres sur IG/TikTok au moment de la rédaction (à valider).")
para("**MESA**")
bullet("« Table » en espagnol/portugais — l'objet universel autour duquel on se rassemble.")
bullet("Ultra court, premium, risque de collision avec Mesa (dev platform).")
para("**TERRANOVA**")
bullet("« Terre nouvelle » — posture d'exploration culinaire.")
bullet("Plus littéraire, meilleur pour le storytelling, moins punchy en app store.")

h2("Slogan")
quote("Cook the world. Shop next door.", PRIMARY)
para("FR : « Cuisine le monde. Trouve-le à côté. »")
para("Fonctionne parce que : verbes d'action, opposition globale/locale, lisible en 1,5 seconde sur une storefront.")

h2("Positionnement")
para("**Premium accessible.** Pas un blog de recettes, pas une appli d'itinéraires, pas une box — un tiers-lieu digital entre les trois. Culture + food + tech, avec une exigence d'authenticité portée par la communauté et validée éditorialement.")

h2("Univers de marque")
bullet("**Ton :** chaleureux, précis, sans mise à distance ironique. On parle comme un ami qui a voyagé et qui cuisine bien — jamais comme un critique gastronomique.")
bullet("**Valeurs :** authenticité, accessibilité, respect des cultures, curiosité, anti-gaspi.")
bullet("**Image perçue :** entre Airbnb (culture + trust), Too Good To Go (anti-gaspi + local) et NYT Cooking (éditorial).")
divider()

# 2. IDENTITÉ VISUELLE
h1("2. Identité visuelle")
h2("Palette de couleurs")
simple_table(
    ["Nom", "HEX", "Rôle"],
    [
        ["Terracotta", "#C85A3A", "Couleur signature — CTA, logo, accents"],
        ["Crème", "#F5EEE0", "Fond principal, respirations"],
        ["Vert safran", "#3E5C3E", "Score authenticité, success states"],
        ["Charcoal", "#1C1B1A", "Texte principal"],
        ["Or doux", "#D9A441", "Highlights premium, badges"],
        ["Gris chaud", "#6B6560", "Texte secondaire, métadonnées"],
    ],
    [1.6, 1.2, 3.7],
)

h2("Typographie")
bullet("**Display :** Fraunces (serif expressif, chaleureux, moderne) — titres, hero recette.")
bullet("**Texte :** Inter (sans-serif, lisibilité mobile exceptionnelle) — UI, corps, listes.")
bullet("**Monospace :** JetBrains Mono — quantités, timers, métadonnées.")

h2("Style visuel")
para("Minimaliste food — plein cadre, lumière naturelle, grain subtil, jamais de filtre saturé. Photo de plat > illustration. Produits et ingrédients shootés en nature morte (bois, céramique, lin). Lifestyle : mains, gestes, quartier — pas de visages studio.")

h2("Direction des images")
bullet("**Plats :** top-shot ou 45°, fond neutre, un seul plat héros par écran.")
bullet("**Ingrédients :** macro, texture visible, étiquettes d'origine lisibles.")
bullet("**Lifestyle :** scènes de marché, mains, quartier — ancrage NYC (Chinatown, Jackson Heights, Sunset Park, Arthur Avenue).")
bullet("**Anti-direction :** pas de flatlay ultra-stylé Pinterest, pas de food photography aseptisée, pas de surjeu ethnique.")
page_break()

# 3. UX
h1("3. UX — parcours utilisateur détaillé")
h2("Onboarding (3 écrans, < 45 secondes)")
num_item("**Écran 1 — Accueil** : logo + slogan + CTA « Commencer ».")
num_item("**Écran 2 — Préférences alimentaires** : sélection multi (vegan, végétarien, pescétarien, halal, kasher, sans gluten, sans lactose, sans arachide, omnivore). Skip possible.")
num_item("**Écran 3 — Profil culinaire** : 3 sliders (niveau : débutant → confirmé ; budget : $ / $$ / $$$ ; temps dispo en semaine : <20min / 30-60min / aucune limite) + permission géolocalisation contextualisée (« pour te montrer les magasins proches »).")
para("Pas de compte requis à l'onboarding. Le mur d'auth arrive au moment de sauvegarder un favori ou de générer une liste d'achats — friction placée où la valeur est déjà perçue.")

h2("Home")
bullet("**Bandeau « Pour toi »** : 3-5 recettes recommandées par IA (préférences + historique + saison + budget).")
bullet("**Cuisine du jour** : une recette éditoriale grand format, rotation quotidienne, narrée (origine, anecdote, chef ou foyer source).")
bullet("**Par pays** : grille horizontale de cartes (Japon, Mexique, Éthiopie, Liban, Inde, Sénégal…).")
bullet("**Rapide & facile** : filtrées < 30 min.")
bullet("**Budget-friendly** : < $8 par personne.")
bullet("**Tendance dans ton quartier** : recettes populaires chez tes voisins géographiques.")

h2("Recherche")
bullet("Barre universelle (plat, pays, ingrédient, chef, magasin).")
bullet("Filtres : pays d'origine, ingrédient clé, temps, difficulté, régime, budget, authenticité minimum.")
bullet("Recherche vocale + recherche par photo (scan d'un plat au restaurant → recettes similaires).")
bullet("Historique et suggestions intelligentes (« tu as cherché bibimbap — essaie aussi japchae »).")

h2("Page Recette")
num_item("**Hero visuel** plein cadre (image ou vidéo 6s).")
num_item("**Métadonnées** : temps total, difficulté, origine géographique (avec pin carte), **score d'authenticité** (0-100), coût moyen / personne.")
num_item("**Narratif court** : 2-3 lignes sur l'origine, famille/chef source.")
num_item("**Ingrédients** : liste avec toggle portions (1 → 12 personnes, recalcul auto), substitutions possibles signalées.")
num_item("**Étapes** : pas-à-pas, timer intégré, mode cuisine (écran toujours allumé, grosse typo, commande vocale « next »).")
num_item("**CTA flottant** « Voir où acheter » — passe à la Map filtrée sur les ingrédients manquants du garde-manger.")
num_item("**Bas de page** : variations régionales, « qui a cuisiné ça cette semaine », commentaires communauté.")

h2("Map")
bullet("Magasins géolocalisés, pins colorés selon score d'authenticité (vert → rouge).")
bullet("Filtres : distance, fourchette de prix, score auth minimum, ouvert maintenant.")
bullet("Fiche magasin : photos, horaires, spécialités, avis communautaires, « confirmé par [X] utilisateurs qui ont trouvé cet ingrédient ».")
bullet("Itinéraire optimisé multi-magasins (algorithme TSP simplifié sur ~5 points max).")

h2("Liste d'achats")
bullet("Génération auto depuis une ou plusieurs recettes, fusion des ingrédients communs.")
bullet("Regroupement intelligent par magasin suggéré (optimisation prix + distance + one-stop quand possible).")
bullet("Quantités ajustables, unités convertibles (imperial ↔ métrique).")
bullet("Check en magasin (swipe pour valider), partage liste (coloc, famille, couple).")
bullet("Export (Apple Reminders, Google Keep, Notion via share sheet).")
page_break()

# 4. FEATURES
h1("4. Features différenciantes")
simple_table(
    ["Feature", "Ce que ça fait", "Pourquoi ça change tout"],
    [
        ["Score d'authenticité", "Note 0-100 par recette, calculée à partir de sources (chefs natifs, communauté de la diaspora, validation éditoriale).", "Seul produit sur le marché qui dit explicitement si une recette est « la vraie ». Différenciation mémorable et citable."],
        ["Route optimizer", "Itinéraire multi-magasins optimal depuis ta position, tenant compte prix + distance + authenticité.", "Transforme une corvée (3-4 magasins) en parcours plaisant. Moment « wow » à l'usage."],
        ["Scan recette → ingrédients", "Photo d'une recette papier / screenshot blog → parsing IA vers ingrédients + map.", "Pont entre l'existant (livres, blogs, stories) et l'achat local. Réduit la friction d'entrée."],
        ["Budget mode", "Active un filtre global : recettes, substitutions, magasins, quantités ajustent pour rester sous $X/semaine.", "Répond à la contrainte #1 des jeunes actifs urbains et expats."],
        ["Pantry AI", "Inventaire du garde-manger (scan ticket / manuel) → suggestions de recettes utilisant ce que tu as déjà.", "Réduit le gaspillage, augmente la rétention (ouvre l'app même sans intention de cuisiner un plat précis)."],
    ],
    [1.3, 2.4, 2.8],
)
page_break()

# 5. STRUCTURE
h1("5. Structure de l'app (arborescence)")
para("Navigation principale : tab bar 5 entrées. Profondeur max 3 niveaux pour tout parcours critique.")
bullet("**Home** — feed personnalisé, éditorial, découverte.")
bullet("**Search** — barre + filtres + recherche vocale/photo.")
bullet("**Recipe** (stack accessible depuis Home, Search, Map, Shopping list) — hero, ingrédients, étapes, CTA shop.")
bullet("**Map** — carte magasins + fiche magasin + itinéraire.")
bullet("**Shopping list** — listes actives, historique, partage.")
bullet("**Profile** — préférences, pantry, favoris, historique, paramètres, premium.")
para("Écrans système : onboarding, paywall premium, invitation / share, auth (login, signup, magic link).")
page_break()

# 6. ECONOMIE
h1("6. Modèle économique")
simple_table(
    ["Source", "Mécanisme", "Comment ça génère du revenu"],
    [
        ["Sponsoring magasins", "Placement premium dans la map + badge « partenaire certifié » + fiche enrichie.", "Fee mensuel fixe ($80–$300/mois selon taille magasin) + upsell CPM pour mises en avant contextuelles (« magasin suggéré » sur une recette matchante)."],
        ["Affiliation", "Partenariats avec Instacart, Fresh Direct, Amazon Fresh, Weee!, Umamicart, HMart online.", "Commission 3–8% sur panier référé. User reste dans Savora pour la découverte, checkout délégué."],
        ["Publicité native", "Brand content de marques food internationales (Kikkoman, Goya, Maggi, marques challengers diasporiques).", "CPM premium ($15–$40) grâce au ciblage ultra-pertinent (ex : une marque coréenne sponsorise des recettes coréennes)."],
        ["Premium utilisateur", "Abonnement $4,99/mois ou $39/an.", "No ads, scan illimité, Pantry AI, recettes exclusives chef, export liste avancé, mode offline, pré-commande chez partenaires."],
        ["Données agrégées (long terme)", "Insights anonymisés sur tendances ingrédient par quartier.", "Vendu à marques CPG (sous RGPD/CCPA, opt-in clair). Marge très élevée mais dépend d'un volume d'usage significatif."],
    ],
    [1.5, 2.5, 2.5],
)
para("**Unit economics cibles année 2 :** ARPU mixte ~$2,40/mois, CAC cible < $4 via contenu organique + referral, LTV > $30 sur 18 mois. Break-even atteignable autour de 80-120k MAU à NYC.")
page_break()

# 7. CONCURRENCE
h1("7. Analyse concurrentielle")
simple_table(
    ["Acteur", "Ce qu'il fait bien", "Sa limite pour notre cas", "Où Savora gagne"],
    [
        ["Google Maps", "Index exhaustif de magasins, reviews, itinéraires.", "Ne sait pas quel ingrédient est vendu, pas de score d'authenticité, pas de lien recette.", "Recette → ingrédient → magasin en 1 flow."],
        ["Yelp", "Avis riches, photos, fiches profondes.", "Orienté restaurants, pas « je cuisine chez moi ». Pas de layer recette/ingrédient.", "Intention « cook at home » couverte de bout en bout."],
        ["HelloFresh / Blue Apron", "Box clé-en-main, planification.", "Fermé, cher, peu authentique, pas d'agencement local, peu de diversité cuisines du monde.", "Ouvert, local, authentique, 10× moins cher, choix illimité."],
        ["NYT Cooking / Kitchen Stories", "Qualité éditoriale, recettes.", "Pas de lien vers achat local, ingrédients rares introuvables.", "Pont entre recette et disponibilité réelle."],
        ["Whisk / Paprika", "Liste de courses générée depuis recettes.", "Pas d'intelligence géo ni authenticité.", "Liste optimisée multi-magasins + score auth."],
    ],
    [1.2, 1.9, 1.9, 1.5],
)
para("**Moat défendable :** (1) le graphe ingrédient ↔ magasin ↔ authenticité est coûteux à construire et se renforce avec l'usage (data network effect), (2) la communauté diasporique qui valide l'authenticité est une asset humaine reproductible difficilement, (3) les partenariats magasins exclusifs deviennent des accords multi-villes.")
page_break()

# 8. MVP
h1("8. MVP — lançable en 30 jours")
h2("Scope features")
bullet("50-80 recettes curated, 10 pays × 5-8 plats (Japon, Corée, Chine, Inde, Mexique, Pérou, Liban, Maroc, Éthiopie, Italie).")
bullet("Recherche par plat / pays / ingrédient (full-text Postgres, pas besoin d'Algolia en V0).")
bullet("Page recette complète (hero, métadonnées, ingrédients ajustables, étapes).")
bullet("Map avec 80-120 magasins NYC (5 quartiers clés : Chinatown Manhattan, Flushing, Jackson Heights, Sunset Park, Arthur Avenue).")
bullet("Liste d'achats basique (sans optimisation multi-magasins au démarrage).")
bullet("Comptes utilisateur, favoris, préférences alimentaires.")
bullet("Score d'authenticité hardcodé sur les 80 recettes initiales (version éditoriale).")
para("**Hors scope V0 (parking) :** Pantry AI, scan recette, route optimizer, mode offline, paywall, social feed. On vendra l'abonnement après avoir prouvé la rétention organique.")

h2("Tech stack")
simple_table(
    ["Couche", "Choix", "Pourquoi"],
    [
        ["Mobile", "React Native (Expo, TypeScript)", "Un seul codebase iOS/Android, OTA updates, écosystème mature."],
        ["Backend / DB", "Supabase (Postgres, Auth, Storage, Edge Functions)", "Lance en quelques heures, Row Level Security robuste, scale jusqu'à 500k utilisateurs sans refactor."],
        ["Map", "Mapbox GL Native", "Meilleure customisation que Google Maps, pricing raisonnable jusqu'à 50k MAU."],
        ["Recherche", "Postgres full-text → Algolia en V1", "Start simple, migre quand pertinent."],
        ["Analytics", "PostHog self-hosted", "Open source, funnels/session replay, pas de vendor lock-in."],
        ["Notifs", "Expo Push + OneSignal", "Gratuit, fiable."],
        ["IA (V1)", "OpenAI API (parsing recette, recos)", "Pay-as-you-go, wrapping léger."],
        ["Paiement (V1)", "RevenueCat + Stripe", "Gestion des abonnements cross-platform sans refaire la logique."],
    ],
    [1.4, 2.1, 3.0],
)

h2("Données nécessaires")
bullet("**Recettes** : rédaction interne + partenariats blogueurs/créateurs diasporiques (licensing ou partage de revenus).")
bullet("**Magasins** : scraping Google Places (API officielle) + enrichissement manuel + vérification terrain (~40h de travail pour 100 magasins).")
bullet("**Mapping ingrédient ↔ magasin** : base manuelle initiale (matrice 300 ingrédients × 100 magasins = 30k cellules, partiellement pré-remplie avec heuristique). Crowdsourcé par la communauté ensuite (un user signale « j'ai trouvé ça ici » → +1 à la data).")
bullet("**Score d'authenticité V0** : grille éditoriale (source de la recette, validation par un native speaker, régionalité précisée).")

h2("Équipe MVP")
bullet("1 fondateur produit/ops (toi) — ops, data, partenariats, contenu.")
bullet("1 dev full-stack React Native/Supabase (contractor ou CTO).")
bullet("1 éditorialiste/community manager part-time (recettes, authenticité, réseaux sociaux).")
bullet("Budget MVP 30 jours : $8–15k hors temps fondateur (dev contractor + Supabase/Mapbox/assets + tournage photos initiales).")
page_break()

# 9. UI DESIGN
h1("9. UI design")
h2("Principes")
bullet("**Mobile-first absolu.** Thumb zone sacrée : toute action critique atteignable au pouce droit.")
bullet("**Hiérarchie visuelle** : image > titre > méta > action. Un seul CTA primaire par écran.")
bullet("**Respiration** : marges généreuses, fond crème, typographie confortable (corps 16-17pt).")
bullet("**Micro-interactions sobres** : transitions < 300ms, haptic feedback sur CTA critiques, jamais d'animation gratuite.")

h2("Écrans principaux")
para("**Home** — status bar transparente, header discret, cards grandes (ratio 16:10) empilées verticalement, tab bar opaque bas. Navigation latérale par sections horizontales scrollables.")
para("**Recipe** — hero 60% hauteur écran, metadata row en flottant sur dégradé foncé, contenu en sheet qui remonte. CTA « Shop these ingredients » en bottom bar persistant.")
para("**Map** — map plein écran, bottom sheet 30% par défaut, extensible à 80% (liste détaillée). Filtres en chips scrollables haut d'écran. Pin magasin sélectionné : fiche se déploie.")
para("**Shopping list** — regroupement par magasin avec header contextuel (nom, distance, prix estimé total), swipe gauche pour check, swipe droit pour supprimer, pull-to-refresh pour ré-optimiser.")
para("**Profile** — liste minimale, premium highlighted, accès rapide à pantry et favoris.")
page_break()

# 10. COPYWRITING
h1("10. Copywriting")
h2("App Store")
para("**Titre :** Savora — Cook Global, Shop Local")
para("**Sous-titre :** World recipes + nearby shops")
h2("Description (≈ 180 mots)")
para("Savora te fait voyager sans quitter ta cuisine. Bibimbap, cochinita pibil, injera, tagine, pho, arepas — des recettes authentiques validées par des chefs et des cuisiniers natifs, avec une seule promesse : tu sais exactement où acheter chaque ingrédient autour de toi.")
para("Choisis un plat. Savora te montre les ingrédients, leur disponibilité dans les magasins proches, un score d'authenticité par recette, et un itinéraire optimisé pour faire tes courses en un seul aller-retour. Ajuste les quantités, remplace un ingrédient, ou active le mode budget — l'app s'adapte.")
para("Conçue à New York, Savora couvre plus de 40 cuisines du monde et référence des centaines de commerces spécialisés. Expats, jeunes actifs, curieux ou passionnés : l'app qui te rend capable de cuisiner (vraiment) le monde entier, à côté de chez toi.")
para("Gratuit. Sans pub intrusive. Premium optionnel pour scan illimité, garde-manger intelligent et recettes exclusives.")

h2("CTA principaux (banque)")
bullet("« Cook this tonight »")
bullet("« Find these ingredients »")
bullet("« See on the map »")
bullet("« Start your list »")
bullet("« Swap an ingredient »")
bullet("« Optimize my route »")
page_break()

# 11. LANCEMENT
h1("11. Stratégie de lancement")
h2("Contenu — TikTok & Instagram")
bullet("**Série pilier** « Where to find [rare ingredient] in NYC » — 60s, format interview express chez le marchand, 3 vidéos/semaine. Objectif : devenir le réflexe de recherche pour les food curious.")
bullet("**Série « Cook the world in 60s »** — format recette express avec overlay magasin NYC où acheter l'ingrédient clé.")
bullet("**Street talks** — micro-trottoirs dans Chinatown, Jackson Heights, Sunset Park. Format court + feed documentaire long (IG Reels + YouTube Shorts).")
bullet("**UGC** — hashtag #CookWithSavora, repost systématique des premiers users.")

h2("SEO (recettes + local)")
bullet("**Pages recettes** indexables (structured data Recipe schema) — chaque recette = une landing page publique.")
bullet("**Pages locales** : « Best Korean grocery stores in Brooklyn », « Where to buy injera in NYC », « The 12 best Mexican markets in Queens » — long-tail high intent.")
bullet("Partenariats backlinks avec blogs food NYC (Eater, The Infatuation, Serious Eats via pitches contributions).")

h2("Acquisition sans budget")
bullet("**Communautés existantes** : groupes Facebook expats (Koreans in NYC, French in New York, Mexicanos en NY, Indians in NYC…), subreddits r/AskNYC, r/FoodNYC, r/Cooking.")
bullet("**Ambassadeurs natifs** : recrute 20-30 cuisiniers diasporiques qui valident les recettes de leur culture en échange de visibilité + premium à vie + paiement symbolique.")
bullet("**Partenariats magasins** : programme « Savora Verified » gratuit pour les 50 premiers magasins, qui reçoivent un sticker vitrine + QR vers leur fiche.")
bullet("**Referrals** : chaque user qui invite 2 amis débloque 1 mois premium + une recette signature exclusive.")
bullet("**Presse** : pitch early à Grub Street, Eater NY, The City, Curbed — angle « a love letter to NYC's immigrant food scene ».")
page_break()

# 12. BONUS
h1("12. Bonus")
h2("3 idées virales")
num_item("**« Authenticity test »** — un user soumet une recette, un cuisinier natif de la culture la note en vidéo (format split-screen duo TikTok). Conflit productif, haute rétention, contenu inépuisable.")
num_item("**« Cook My Heritage »** — user raconte sa recette de famille en vidéo 60s, l'app la référence avec son histoire, elle devient publique. Émotionnel, viral, et construit un moat contenu irréplicable.")
num_item("**Filter IG/Snap « Ingredient Hunter »** — filtre AR qui montre un ingrédient rare (ex : feuille de bananier) et demande « trouve-le dans les 24h à NYC ». Tagger le magasin donne des points.")

h2("Hack de croissance")
para("**Programme « Diaspora Ambassadors » gamifié.** 30-50 cuisiniers de diasporas recrutés à la main, chacun devient référent pour sa culture. Ils gagnent un % sur chaque user qui valide une recette qu'ils ont ambassadeur-validée (type creator fund). Impact : (1) ils recrutent leur propre communauté, (2) l'authenticité devient un asset humain mesurable, (3) storytelling presse naturel.")

h2("Feature future game changer")
para("**« Pantry Instant » + planning semaine IA.** L'utilisateur scanne ses tickets de caisse (email receipts ou photo). L'app reconstitue son garde-manger en temps réel, connaît ses préférences, son budget et son emploi du temps, et génère automatiquement un planning de la semaine + la liste d'appoint minimale. C'est la version « Netflix for what to cook tonight » — le moment où Savora passe d'outil de découverte à copilote quotidien indispensable.")
page_break()

# =====================================================
# VERSION 2 — AGRESSIVE CROISSANCE / VIRALITÉ
# =====================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VERSION 2")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MUTED
p.paragraph_format.space_before = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Agressive croissance / viralité")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = POP_RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Social-first · content-as-product · K-factor obsessed")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(18)
divider()

# 1. BRANDING v2
h1("1. Branding", POP_RED)
h2("Nom de marque — 3 propositions")
para("**MAMA MAP** (recommandé)")
bullet("Émotion + objet : tout le monde a une mama/abuela/halmoni/jiddo/nonna qui cuisine bien.")
bullet("Visuel immédiat : map = localisation, promesse lisible en 0,5s.")
bullet("Ultra-viral en hashtag : #MamaMap, #MamaMapNYC, #AskMamaMap.")
para("**YUMMIGO**")
bullet("Yummy + Go → cuisine + déplacement. Sonorité enfantine mémorable.")
bullet("Moins premium, très TikTok-friendly.")
para("**COOKALA**")
bullet("Cook + « alla » (comme en italien). Rythme, exotisme, nom de danse.")
bullet("Excellent comme chant/jingle (« cookalaaa ! »).")

h2("Slogan")
quote("Cook like your grandma. Shop like a local.", POP_RED)
para("FR : « Cuisine comme ta grand-mère. Achète comme un local. »")
para("Fonctionne parce que : déclenche immédiatement une image mentale + une mini-histoire. Zéro besoin d'explication.")

h2("Positionnement")
para("**Social-first, community-owned, content-as-product.** Pas une app food — un mouvement culturel autour de la cuisine du monde à NYC. Le produit est le vecteur, la communauté est la valeur, le contenu est le moteur d'acquisition.")

h2("Univers de marque")
bullet("**Ton :** drôle, sincère, un peu chaotique, jamais condescendant. Voix de la cousine qui ramène toute sa culture à table.")
bullet("**Valeurs :** fierté culturelle, chaos joyeux, anti-cringe, communauté > algorithme.")
bullet("**Image perçue :** entre Duolingo (mascotte attachante, tone of voice clivant), Cash App (pop/urban) et Airbnb Stories (émotion + authenticité).")
divider()

# 2. IDENTITÉ VISUELLE v2
h1("2. Identité visuelle", POP_RED)
h2("Palette de couleurs")
simple_table(
    ["Nom", "HEX", "Rôle"],
    [
        ["Tomate", "#FF4E3A", "Signature, CTA, logo"],
        ["Turmeric", "#FFC933", "Accents, alertes joyeuses, badges"],
        ["Noir profond", "#111111", "Texte principal, blocs poppy"],
        ["Blanc chaud", "#FFFBF3", "Fond principal"],
        ["Menthe", "#4FD1A8", "Success / ingrédient trouvé"],
        ["Magenta", "#E63E7A", "Highlights, pastilles challenge"],
    ],
    [1.6, 1.2, 3.7],
    header_color="FF4E3A",
)

h2("Typographie")
bullet("**Display :** Chillax ou GT Maru (arrondi, pop, très TikTok-friendly).")
bullet("**Texte :** Manrope (sans-serif moderne, lisible, gratuit).")
bullet("**Accent :** Caveat ou Handlee (écriture manuscrite) — pour les citations, annotations, note de mama.")

h2("Style visuel")
para("Pop food + DIY aesthetic. Plats shootés au flash direct (vibe 2000s/iPhone raw), pas de lumière studio aseptisée. Emojis intégrés (mesurés), stickers, typographie qui déborde, cadres colorés façon Polaroid. On assume le « trop », on refuse le « trop propre ». Inspiration : TikTok feed, Depop, Duolingo, Sleepy Dog.")

h2("Direction des images")
bullet("**Plats :** flash direct, assiette en main, cadre inclus volontairement.")
bullet("**Ingrédients :** zoom extrême, texture, parfois sous-titré en manuscrit (« my mom uses this »).")
bullet("**Lifestyle :** visages, mains, rires, chaos de cuisine. UGC roi.")
bullet("**Anti-direction :** pas de food photography classique, pas de minimalisme premium, pas de tonalité sérieuse.")
page_break()

# 3. UX v2
h1("3. UX — parcours utilisateur (version social-first)", POP_RED)
h2("Onboarding (3 écrans, gamifié)")
num_item("**Écran 1 — Quiz « Tell us who you cook like »** — 5 questions pop (ex : « Ta mama ajoute plus d'ail ou plus de gingembre ? », « Ton plat d'urgence à 2h du matin ? »). Chaque réponse donne un badge starter.")
num_item("**Écran 2 — Permission géoloc** avec visuel « Unlock your hood's food map » + animation map qui s'allume.")
num_item("**Écran 3 — Invite 2 friends → skip the waitlist** (mécanique Clubhouse/Superhuman). Résistance acceptable car récompense claire.")

h2("Home — feed vertical swipeable (type TikTok)")
bullet("**Feed principal** : recettes en vidéo 15-45s, plein écran, swipe up pour la suivante. Chaque vidéo : bouton « Cook this » (passe en mode recette) + « Shop this » (map instantanée).")
bullet("**Double-tap** pour sauver, **hold** pour ajouter au planning.")
bullet("**Tiles surimposés** : ingrédients épinglés, tag magasin, timer « trouvable en 7 min de chez toi ».")
bullet("**Stories horizontales top** : challenges du jour, lives communauté, drops de quartier.")

h2("Recherche")
bullet("Par mood (« late night », « comfort », « show off », « heal me », « hangover »), par hashtag, par quartier, par créateur.")
bullet("Recherche vocale + recherche par photo (scan d'un plat au restaurant).")

h2("Page Recette — dense mais ludique")
bullet("Vidéo loop hero, créateur taggé, commentaires ouverts type TikTok.")
bullet("Score d'authenticité visible en gros + qui l'a validé (visages, liens profils).")
bullet("Ingrédients : chaque ligne cliquable → ouvre une mini-map du magasin le plus proche avec itinéraire direct.")
bullet("Étapes en carrousel stories, chacune 5-10s, voix du créateur si dispo.")
bullet("**CTA triple bottom** : Shop · Cook · Post my version.")

h2("Map — social")
bullet("Magasins avec activité en direct : « 12 users came here this week », « $$ spent here average ».")
bullet("Posts communauté épinglés (un user a trouvé un ingrédient rare → sticker sur le magasin).")
bullet("Mode AR : pointe ton téléphone dans la rue, les magasins « quest » apparaissent en overlay.")

h2("Liste d'achats — ludique")
bullet("« Quest mode » : liste devient mini-aventure avec itinéraire, checkpoints, temps estimé, points gagnés.")
bullet("Partage 1-tap avec qui cuisine avec toi (statut en direct des items cochés).")
bullet("Streak bonus : 7 jours de cook-at-home = débloque recette exclusive.")
page_break()

# 4. FEATURES v2
h1("4. Features différenciantes — version viralité", POP_RED)
simple_table(
    ["Feature", "Ce que ça fait", "Pourquoi ça devient viral"],
    [
        ["Feed vidéo swipeable", "TikTok-style feed de recettes 15-45s, chaque vidéo shoppable.", "Rétention type social, pas type utilitaire. Users ouvrent sans intention précise."],
        ["Grandma Mode", "Upload audio ou vidéo d'un proche qui te guide pendant la cuisine. L'app sync avec les étapes.", "Émotion pure. Shareable. Irremplaçable. Crée un moat de contenu familial."],
        ["Hood Chef Battles", "Classement hebdo par quartier : qui cuisine le plus, avec quel score auth, avec quel impact communautaire.", "Rivalité locale saine, contenu natif pour TikTok, FOMO de quartier."],
        ["Mystery Ingredient Drop", "Chaque lundi, un ingrédient rare est « droppé » dans NYC. Premiers à le trouver + l'utiliser gagnent des récompenses (bon resto, premium, drops).", "Event récurrent = rétention. Chasse = contenu UGC gratuit."],
        ["Cook & Post", "Caméra in-app, filtres, tags magasins pré-remplis. Post direct depuis l'étape finale.", "L'acte de cuisiner devient l'acte de créer. UGC = moteur d'acquisition."],
    ],
    [1.5, 2.5, 2.5],
    header_color="FF4E3A",
)
page_break()

# 5. STRUCTURE v2
h1("5. Structure de l'app — version social", POP_RED)
para("Tab bar 5 entrées, centrée sur le social et la création.")
bullet("**Feed** — home social, recettes vidéo, stories, drops.")
bullet("**Map** — magasins, quartiers, communauté live.")
bullet("**Create** (bouton central proéminent) — Cook & Post, upload Grandma Mode, challenge entry.")
bullet("**Lists** — quests, planning, pantry.")
bullet("**Profile** — streaks, badges, leaderboard quartier, favoris, friends.")
para("Écrans système : onboarding gamifié, waitlist FOMO, referral tree, paywall, auth.")
page_break()

# 6. ECONOMIE v2
h1("6. Modèle économique — monétisation agressive", POP_RED)
simple_table(
    ["Source", "Mécanisme", "Comment ça génère du revenu"],
    [
        ["Sponsoring magasins", "Drops sponsorisés hebdomadaires, fiche pro, badges spéciaux.", "$50–$500/drop + abonnement pro mensuel."],
        ["Affiliation", "Delivery partners + brands DTC food.", "Commission 5–12% sur paniers, plus élevée qu'en V1 grâce au volume viral."],
        ["Brand challenges sponsorisés", "Une marque finance un challenge (ex : Kikkoman × #UmamiChallenge) avec prize pool.", "Tickets $10k–$100k par brand challenge. Forte marge, bundle créator+audience."],
        ["Creator fund partagé", "Top 500 creators de l'app reçoivent un % des revenus ads/premium liés à leurs recettes.", "Pas un coût mais un investissement en rétention créateurs. Alignement d'intérêts."],
        ["Premium utilisateur", "$6,99/mois — mais vendu comme « VIP pass » avec features cosmétiques (badges, drops early access, resto partenaires).", "Tarif plus élevé accepté grâce au statut social."],
        ["Marketplace pro", "Petits commerces non-tech peuvent s'onboarder en 2 min, payer $29/mois pour fiche premium + analytics.", "SaaS micro-business, haute rétention (ancre leur visibilité locale)."],
    ],
    [1.5, 2.5, 2.5],
    header_color="FF4E3A",
)
para("**Hypothèse de scale :** K-factor visé ≥ 1,15 grâce aux mécaniques Cook & Post + referral waitlist + challenges. ARPU plus bas que V1 mais volume 4-6× supérieur si la viralité tient. Valorisation plus agressive en série A grâce au narratif social-media × food.")
page_break()

# 7. CONCURRENCE v2
h1("7. Analyse concurrentielle — version viralité", POP_RED)
simple_table(
    ["Acteur", "Ce qu'il fait bien", "Sa limite", "Où Mama Map gagne"],
    [
        ["TikTok Food", "Viralité, découverte, créateurs.", "Aucun lien avec achat local, aucun score auth, contenu éphémère.", "Contenu shoppable ancré géo + archivable + structuré."],
        ["Google Maps", "Index magasins.", "Zéro social, zéro recette, zéro contenu.", "Couche sociale + contenu vidéo natif."],
        ["Yelp", "Reviews.", "Orienté resto, vieillissant, pas social.", "Génération Z-native, recettes au cœur, cook-not-eat."],
        ["HelloFresh", "Box.", "Fermé, cher, non local.", "Ouvert, gratuit, viral, créateur-driven."],
        ["Reddit r/NYCFood, groupes FB", "Communautés passionnées.", "Désorganisés, pas mobiles, pas monétisés.", "Centralise l'énergie dans un produit structuré."],
    ],
    [1.2, 1.9, 1.9, 1.5],
    header_color="FF4E3A",
)
para("**Moat défendable :** (1) le contenu UGC indexé sur la map devient un asset impossible à répliquer, (2) les creators deviennent dépendants de l'audience qu'ils bâtissent sur Mama Map, (3) les drops hebdomadaires créent un rendez-vous que les concurrents ne peuvent pas copier sans communauté existante.")
page_break()

# 8. MVP v2
h1("8. MVP — 30 jours, version virale", POP_RED)
h2("Scope features")
bullet("Feed vertical swipeable avec 40-60 recettes vidéo (produites in-house + 5-8 creators early).")
bullet("Map avec ~80 magasins NYC, focus 3 quartiers iconiques.")
bullet("Cook & Post (caméra + tags magasin + hashtag auto).")
bullet("Waitlist gamifiée avec position qui remonte via invitations.")
bullet("Referral tree visible (tu vois qui a débloqué grâce à toi → dopamine sociale).")
bullet("Challenge hebdomadaire #1 dès semaine 2 post-launch.")
para("**Hors scope V0 :** Pantry AI, Grandma Mode complet (mais prévu V1, très tôt), map AR, marketplace pro.")

h2("Tech stack")
bullet("**Mobile :** React Native (Expo) + Reanimated 3 + Skia pour fluidité feed vidéo.")
bullet("**Vidéo :** Mux (CDN + transcoding) ou Cloudflare Stream — évite de réinventer TikTok.")
bullet("**Backend :** Supabase + Redis cache (Upstash) pour feed & counters.")
bullet("**Map :** Mapbox.")
bullet("**Notifs & engagement :** OneSignal + Customer.io pour lifecycle.")
bullet("**Analytics :** PostHog + Mixpanel (funnels K-factor, cohortes).")
bullet("**Creator backend :** Notion + Airtable au début (pas de dashboard custom avant 100 creators actifs).")

h2("Données initiales")
bullet("**40-60 vidéos** produites en 2 semaines par une équipe de 2 (creator + camera op), tournées dans 3 quartiers.")
bullet("**Base magasins** enrichie manuellement.")
bullet("**Seeding communauté** : 20 creators diasporiques pré-onboardés avant launch (3 semaines de warm-up).")

h2("Équipe MVP")
bullet("1 fondateur produit + community.")
bullet("1 dev full-stack React Native (cross-team avec motion designer part-time).")
bullet("1 video producer/editor (in-house, passé TikTok/agency).")
bullet("1 community manager multilingue.")
bullet("Budget MVP 30 jours : $18–28k hors temps fondateur (production vidéo est le surcoût majeur).")
page_break()

# 9. UI v2
h1("9. UI design — version pop & social", POP_RED)
h2("Principes")
bullet("**Full bleed vidéo** par défaut. L'image gouverne. UI overlay avec glassmorphism léger.")
bullet("**Motion-first** : micro-anims constantes, haptic fort, transitions expressives. L'app doit se sentir vivante.")
bullet("**Hiérarchie** : vidéo > titre > créateur > tags > CTA. CTA en couleur signature, impossible à rater.")
bullet("**Gamification discrète mais présente** : XP, badges, streaks, mais jamais bloquant l'utilité.")

h2("Écrans principaux")
para("**Feed** — full screen vidéo, swipe vertical, double-tap animé (cœur), bouton latéral « Shop » animé en pulse, titre recette + creator dans la lower third.")
para("**Recipe** — vidéo hero en boucle, bouton Shop flash rouge tomate en sticky bottom, étapes en carrousel stories avec haptic à chaque transition.")
para("**Map** — full screen map, pastilles de quartier animées, drops actifs en pulse, bottom sheet avec posts live communauté.")
para("**Create** — bouton central tab bar, ouvre en full screen caméra avec filtres, tags pré-remplis depuis la recette en cours, partage 1-tap.")
para("**Profile** — grille tiktok-like de tes posts, streak + badges top écran, leaderboard quartier accessible.")
page_break()

# 10. COPY v2
h1("10. Copywriting — version viralité", POP_RED)
h2("App Store")
para("**Titre :** Mama Map — Cook global. Shop local.")
para("**Sous-titre :** World recipes tagged on a map.")
h2("Description (≈ 170 mots)")
para("Mama Map, c'est TikTok pour les cuisines du monde — mais chaque recette est reliée aux magasins près de chez toi. Tu swipes, tu kiffes, tu cuisines.")
para("Une recette coréenne te met l'eau à la bouche ? Tap → tu vois où acheter le gochugaru à 7 min à pied. Envie de pozole ce soir ? Mama Map te montre les 3 marchés mexicains où tu trouveras tout, dans l'ordre optimal.")
para("Filme ton plat. Tag ton magasin. Post. Ta cuisine devient du contenu, ton quartier devient une carte, et tu débloques des recettes exclusives à chaque streak.")
para("Challenges hebdo avec des vrais prix. Drops d'ingrédients rares à chasser dans la ville. Leaderboard par quartier. Et des recettes signées par des cuisiniers natifs, validées par les diasporas qui vivent ici.")
para("Cook like your grandma. Shop like a local. Post like a legend.")

h2("CTA principaux (banque)")
bullet("« Cook this »")
bullet("« Shop this »")
bullet("« Post your version »")
bullet("« Join the challenge »")
bullet("« Unlock this drop »")
bullet("« Skip the waitlist »")
page_break()

# 11. LANCEMENT v2
h1("11. Stratégie de lancement — viralité maximale", POP_RED)
h2("Pré-launch — 6 semaines de hype")
bullet("**Waitlist FOMO** gamifiée dès la semaine 1. Chaque invitation = +5 positions. Objectif : 50-100k waitlist avant ouverture.")
bullet("**Sticker campaign guérilla** dans 10 quartiers : Chinatown, Flushing, Jackson Heights, Sunset Park, Arthur Avenue, Astoria, Washington Heights, Bushwick, Bed-Stuy, LES. Sticker QR qui fait gagner une place.")
bullet("**Teaser content** : série TikTok « Things your bodega doesn't tell you », « Secret ingredients in X neighborhood ».")

h2("Launch — semaines 1-4")
bullet("**#MamaMapChallenge** : cuisine la recette de ta mama, tag son pays d'origine, on géo-référence. Prix : $10k de food credits répartis sur les 50 meilleures vidéos.")
bullet("**Street interviews** virales : « Show me what's in your fridge » + « Where does this come from in NYC ». 5-10 vidéos/semaine.")
bullet("**Nano-influenceurs natifs** : 80-120 comptes entre 1k-10k abonnés, chacun est « ambassadeur d'une cuisine ». Paiement en premium à vie + revenue share + co-création de drops.")
bullet("**Pop-up market day** : partenariat 5 magasins, journée « Mama Map Day » avec dégustations, recettes imprimées et offres exclusives.")

h2("SEO & earned media")
bullet("**Pages « Best X in [neighborhood] »** rédigées par creators natifs, indexables, shoppables.")
bullet("**Pitch presse** : Grub Street, Eater, Mission Local, The Infatuation, TikTok newsletter. Angle « the app that gamified NYC's immigrant food scene ».")
bullet("**Collaborations culturelles** : museums (MoCA, El Museo del Barrio), events (Smorgasburg), food weeks.")
page_break()

# 12. BONUS v2
h1("12. Bonus — max viralité", POP_RED)
h2("3 idées virales")
num_item("**Grandma Mode audio** : les utilisateurs uploadent leur mama/abuela/nonna qui guide une recette en voice. Format extrêmement partageable sur TikTok, émotion maximale, génère du contenu pour la fête des mères / Thanksgiving / Lunar NY. Moat contenu familial.")
num_item("**Hood Chef Battles** : tous les vendredis, classement par quartier (Chinatown vs Flushing vs Sunset Park…). Contenu native pour recap TikTok du lundi, rivalité joyeuse, FOMO hebdo.")
num_item("**Mystery Ingredient Hunt** : chaque lundi, un ingrédient rare est « caché » à NYC. Première personne à le trouver + cuisiner + poster = gagne un dîner chez un resto iconique du quartier. Crée un event hebdomadaire qui ramène les users.")

h2("Hack de croissance")
para("**Waitlist gamifiée façon Superhuman/Clubhouse, combinée à un « referral tree visualisé ».** Tu vois en live l'arbre de tes invitations (qui a débloqué grâce à toi, puis qui ils ont invité à leur tour). Dopamine sociale brutale, boucle virale naturelle. Objectif : 1 user = 3 inscrits. Atteindre 100k waitlist en < 8 semaines est réaliste avec le sticker campaign + les nano-ambassadeurs.")

h2("Feature future game changer")
para("**AI Chef Personnel** : un agent qui connaît ta culture, tes goûts, ton budget, ton planning, ton garde-manger et le contenu communautaire. Il propose chaque matin : « ce soir, tu pourrais faire ceci — tu as 80% des ingrédients, 20% à chercher à 5 min à pied, 12$ total, la recette vient de [creator], déjà cuisinée par 3 de tes amis ». Quand l'app passe de « place où je swipe » à « copilote quotidien de ma vie alimentaire », la rétention devient structurelle et le LTV explose. C'est le moment où Mama Map cesse d'être une app food pour devenir un OS de la vie quotidienne.")
page_break()

# =====================================================
# INVESTOR PITCH
# =====================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OPTION AVANCÉE")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MUTED
p.paragraph_format.space_before = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pitch fondateur — lever des fonds")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = DARK
p.paragraph_format.space_after = Pt(12)

quote("Agis comme un fondateur de startup qui doit lever des fonds et rendre ce projet irrésistible pour des investisseurs.", GOLD)
divider()

h2("Problem")
para("37% des New-Yorkais sont nés hors des US. 800+ langues parlées. Des dizaines de milliers de commerces alimentaires spécialisés. Et pourtant : zéro produit mobile qui relie une recette authentique à l'endroit exact où acheter ses ingrédients. Google Maps n'a pas le contexte recette. Yelp n'a pas le contexte cook-at-home. HelloFresh n'a pas le contexte local ni culturel. Les blogs n'ont pas le contexte géographique. Le marché est fragmenté verticalement — nous agrégeons.")

h2("Solution")
para("Savora / Mama Map est le premier graphe recette × ingrédient × magasin × authenticité, servi mobile-first. Tu choisis ton plat, tu vois où acheter, tu pars faire tes courses en un trajet optimisé. Le score d'authenticité validé par les diasporas nous donne un moat culturel défendable. Le graphe enrichi par les users crée un data network effect classique.")

h2("Market")
bullet("**TAM global :** > $1,5T (global grocery + prepared food discovery, Euromonitor 2025).")
bullet("**SAM addressable (urban multicultural markets, top 20 villes) :** $140B.")
bullet("**SOM 3 ans (NYC + LA + Miami + Toronto + Londres) :** $4,8B.")
bullet("**TAM étroit prouvable** côté recettes + delivery + CPG ads sur cuisines du monde à NYC : > $400M.")

h2("Why now")
bullet("L'essor du cook-at-home post-2020 est durable (≈ +18% vs pré-pandémie en dépenses épicerie US).")
bullet("Gen Z et Millennials expats recherchent activement authenticité culturelle dans leur alimentation.")
bullet("Coût du cooking content en vidéo divisé par 5 en 4 ans grâce à iPhone + outils d'édition + creators.")
bullet("Infrastructure IA (LLMs, vision) rend le parsing recette et le mapping ingrédient ↔ magasin économiquement viable pour la première fois.")

h2("Traction cible")
simple_table(
    ["Horizon", "KPI leading", "KPI lagging"],
    [
        ["Mois 1-3 (MVP NYC)", "20k MAU, K-factor 0,8+", "Rétention D30 ≥ 25%"],
        ["Mois 4-9 (Expand NYC)", "120k MAU, 8k paid", "ARPU $2, CAC < $4"],
        ["Mois 10-18 (2 villes)", "450k MAU, 40k paid", "LTV > $30, CAC < LTV/3"],
        ["Mois 19-36 (5 villes)", "2M MAU, 200k paid", "Contribution margin > 60%"],
    ],
    [2.1, 2.2, 2.2],
    header_color="D9A441",
)

h2("Business model")
para("Mix pluri-sources (sponsoring magasins + affiliation + ads natives + premium + data anonymisée). Point critique : le mix est construit pour que chaque nouvelle ville contribue positivement dès le mois 4 post-launch grâce aux deals magasins et affiliation pré-signés.")

h2("Moat")
bullet("**Data network effect** : chaque user enrichit le graphe ingrédient ↔ magasin.")
bullet("**Moat culturel** : validation diasporique coûteuse à reproduire, renforcé par relations humaines.")
bullet("**Partenariats magasins exclusifs** : onboarding en avance crée une stickiness contractuelle.")
bullet("**Contenu UGC + éditorial** : asset médiatique qui renforce le SEO et l'acquisition organique.")

h2("Ask")
para("**Seed round : $2,2M.** Usage : (1) 45% produit et engineering — équipe de 5 pour 18 mois, (2) 25% content et community (creators, production vidéo, éditorial), (3) 20% growth et paid experimentation à partir du mois 6, (4) 10% buffer et operations. Objectifs fin de runway : 500k MAU, 2 villes live, métriques unit economics prouvées, prêts pour série A de $8-12M.")

h2("Why us")
para("Équipe qui combine product/ops, tech mobile et obsession culturelle. Accès pré-établi à 30+ ambassadeurs diasporiques pour le seeding authentique. Réseau presse food NYC warm. Capacité à exécuter sur les deux versions (sérieuse ET virale) en parallèle — car le MVP est conçu pour que les mécaniques virales soient activables en V1 si les signaux de traction y poussent.")

divider()
para("**La question à se poser :** dans 5 ans, quand je veux cuisiner un plat du monde à NYC, Londres ou Singapour, est-ce que j'ouvre 4 apps — ou une seule ? Savora / Mama Map est cette seule app.")

# add header/footer last so they apply to all sections
add_header_footer()

# Patch settings.xml zoom element: python-docx writes <w:zoom/> without required attr
settings = doc.settings.element
for zoom in settings.findall(qn('w:zoom')):
    if zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

# Save
import sys
out_path = sys.argv[1] if len(sys.argv) > 1 else "/sessions/dreamy-tender-curie/mnt/outputs/Savora_MamaMap_Concept.docx"
doc.save(out_path)
print("Wrote", out_path)
